import fitz  # PyMuPDF for PDF text extraction
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract


# ---------------------------------------------------------------------------
# Tesseract config
# ---------------------------------------------------------------------------

# Point this to your local Tesseract install if needed.
# On Windows, something like:
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """Extract plain text from a PDF file using PyMuPDF."""
    text_parts = []

    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))

    return "\n".join(text_parts).strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file, including paragraphs, tables, headers and footers."""
    try:
        doc = Document(file_path)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)

        # Headers and footers
        for section in doc.sections:
            header = section.header
            footer = section.footer

            for para in list(header.paragraphs) + list(footer.paragraphs):
                if para.text:
                    parts.append(para.text)

        return "\n".join(parts).strip()

    except Exception as exc:
        # Log and fail gracefully
        print(f"Error extracting text from DOCX ({file_path}): {exc}")
        return ""


# ---------------------------------------------------------------------------
# IMAGES (JPEG/PNG)
# ---------------------------------------------------------------------------

def extract_text_from_jpeg(file_path: str) -> str:
    """Extract text from an image using Tesseract OCR."""
    try:
        image = Image.open(file_path)

        # Basic preprocessing: grayscale, contrast, threshold, median filter
        gray = image.convert("L")
        contrast = ImageEnhance.Contrast(gray).enhance(2)

        threshold = contrast.point(lambda p: 255 if p > 128 else 0)
        processed = threshold.filter(ImageFilter.MedianFilter())

        # OEM 3: default, PSM 6: assume a single uniform block of text
        config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(processed, config=config)

        text = text.strip()
        if not text:
            raise ValueError("No text found in the image.")

        return text

    except Exception as exc:
        raise ValueError(f"Error processing image '{file_path}': {exc}") from exc


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

def break_text_into_chunks(text: str, chunk_size: int = 300) -> list[str]:
    """Split text into fixed-size chunks (by characters)."""
    text = text or ""
    return [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]
